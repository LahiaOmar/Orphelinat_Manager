from flask import Flask, request
import json
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.shared import Pt

app = Flask(__name__)
 
@app.route('/')
def index():
	return "Flask server"
 
@app.route('/postdata', methods = ['POST'])
def postdata():
    data = request.get_json()
    print(data)
    pic = "img.jpeg"
    dataLeft=[" –العنوان : مستشفى الغساني – فاس", "رقم الانخراط في : ص.و.ض.ج:9005316", " 157N302770 : رقم الحساب البنكي ", "رقــــــم الــــهـاتـف : 05 35 94 35 87 5"]  
    dataRight = ["مؤسسة الرعاية الاجتماعية ", "المركز الاجتماعي للأطفال المهملين", "رخصة رقم : 09/163", "بتاريخ : 11 دجنبر 2014"]
    salary_array = ["التغطية الصحية", "CNSS", "IGR", "تسبيق", "غيابات"]
    MOD = data['mod']
    data = data['data']
    docu = docx.Document()
    lans_scape = ["kid", "benificiaire", "salary", "employee"]
    salaryFields =["كيفية الأداء", "المبلغ الصافي للأداء", "التغطية الصحية", "CNSS", "IGR", "تسبيق", "غيابات", "تعويضات عائلية", "الأجرة الإجمالية", "زيادة في الأجرة", "الأجرة الأساسية", "ب.ت.و", "المهمة", "الاسم الشخصي والعائلي"]
    mod_title = {"kid":"قائمة الأطفال", "benificiaire" : "قائمة المستفيدين", "spent":"النفقات", 
             "income": "جرد المداخيل", "salary":"جرد الأجور" , "equivalence": "سجل الموازنةالعامة",
             "employee":"قائمةالمستخدمون", "equipment":"قائمة القاعات و التجهيزات", "editOutcome":"ارشيف المواد المستهلكة",
             "editIncome":"جرد المواد الوافدة", "product":"قائمة المنتوجات" }
    
    #title od page 
    
    
    
    # Header 
    table = docu.add_table(rows = 0, cols = 3)
    table.alignment =  WD_TABLE_ALIGNMENT.CENTER
    cells = table.add_row().cells
    text_left = ""
    for i in range(len(dataLeft)):
        text_left += ( dataLeft[i] + "\n" )
    p1 = cells[0].add_paragraph(text_left)
  
    cells[0].width = Cm(6)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2 = cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run()
    run.add_picture(pic, width = 1400000, height = 1400000)

    text_right = ""
    for i in range(len(dataRight)):
        text_right += ( dataRight[i] + "\n" )
    p3 = cells[2].add_paragraph(text_right)
    cells[2].width = Cm(5)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # fin header
    # title of table  :

    title = mod_title[MOD]
    para = docu.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    font = run.font
    font.bold = True
    font.size = Pt(40)
    font.rtl = True
    # add break line
    p = docu.add_paragraph()
    run = p.add_run()
    run.add_break()
    # table
    print(data)
    row = len(data)
    col = len(data[0])
    section = docu.sections[0]
    wi = Cm(section.page_width)
    print("wi = ", wi)
    print("lens : ", row, col)
    if MOD != 'salary':
      mat = docu.add_table(rows = row, cols = col)
      mat.alignment = WD_TABLE_ALIGNMENT.CENTER
      mat.autofit = True
      mat.style = "Table Grid"
      for i in range(row):
        cur_row = mat.rows[i].cells
        for j in range(col):
          word = str(data[i][j])
          # li = word.split(' ')
          # li.reverse()
          # word = ' '.join(li)
          p =  cur_row[j].paragraphs[0]
          run = p.add_run(word)
          if i == 0: 
              run.bold = True
              font = run.font
              font.rtl = True
          p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else :
      table = docu.add_table(rows = row, cols = col)
      table.alignment = WD_TABLE_ALIGNMENT.CENTER
      table.autofit = True
      table.style = "Table Grid"
      a = table.cell(0,2)
      b = table.cell(0,6)
      A = a.merge(b)
      for i in range(row):
        row = table.rows[i].cells
        for j in range(col):
          word=""
          if i == 0 and j == 3:
            p = row[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("اقتطاعات")
            run.bold = True
          if i == 0:
            if j < 2 or j > 6:
              word=salaryFields[j]
              a = table.cell(0,j)
              b = table.cell(1,j)
              A = a.merge(b)
              p = row[j].paragraphs[0]
              p.alignment = WD_ALIGN_PARAGRAPH.CENTER
              run = p.add_run(word)
              run.bold = True
          if i == 1:
            if j >=2 and j<=6:
              word = salaryFields[j]
              p = row[j].paragraphs[0]
              p.alignment = WD_ALIGN_PARAGRAPH.CENTER
              run = p.add_run(word)
              run.bold = True
          if i > 1:
            row[j].text = data[i][j]
          
    sections = docu.sections
    
    for section in sections:
      # change orientation to landscape
      if MOD in lans_scape :
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

      section.top_margin = Cm(0)
      section.bottom_margin = Cm(0)
      section.left_margin = Cm(0.5)
      section.right_margin = Cm(0.5)
      
    docu.save("demo.docx")

    # do something with this data variable that contains the data from the node server
    return json.dumps({"ok":"ok"})
 
if __name__ == "__main__":
	app.run(port=5000, debug=True)      